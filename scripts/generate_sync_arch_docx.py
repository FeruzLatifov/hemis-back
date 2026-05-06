#!/usr/bin/env python3
"""
HEMIS-back / Univer Sync Architecture Report — DOCX Generator

Yaratadi: docs/SYNC_ARCHITECTURE_REPORT.docx
Foydalanish: python3 scripts/generate_sync_arch_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "docs" / "SYNC_ARCHITECTURE_REPORT.docx"

# ============================================================
# Helpers
# ============================================================
doc = Document()

# Default font: Calibri 11
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def p(text, bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return para


def code(text, label=None):
    if label:
        lp = doc.add_paragraph()
        lr = lp.add_run(label)
        lr.bold = True
        lr.italic = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
    # Light grey shading
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        for run in hcells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    # rows
    for ri, row in enumerate(rows, start=1):
        rcells = t.rows[ri].cells
        for ci, val in enumerate(row):
            rcells[ci].text = str(val)
            for run in rcells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = w
    doc.add_paragraph()  # spacer


def bullets(items):
    for it in items:
        para = doc.add_paragraph(it, style="List Bullet")
        for run in para.runs:
            run.font.size = Pt(11)


def numbered(items):
    for it in items:
        para = doc.add_paragraph(it, style="List Number")
        for run in para.runs:
            run.font.size = Pt(11)


def hrule():
    p("─" * 80, italic=True, size=8)


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("UNIVER ↔ HEMIS-BACK\nSYNC ARXITEKTURASI")
trun.bold = True
trun.font.size = Pt(28)
trun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = subtitle.add_run("Texnik hisobot — Kafka-first yondashuv tahlili")
srun.italic = True
srun.font.size = Pt(14)
srun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mrun = meta.add_run("Sana: 2026-05-06   |   Status: Proposal   |   Author: Hemis-back team")
mrun.font.size = Pt(10)
mrun.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Abstract
add_h("Qisqacha xulosa", level=2)
p(
    "Ushbu hujjat 224 OTM (Oliy Ta'lim Muassasalari) ↔ Hemis-back orasidagi sync "
    "arxitekturasi tahlilini taqdim etadi. Hozirgi REST-based fire-and-forget yondashuvning "
    "kamchiliklari, 7 ta alternativ yechim baholangan va tanlangan yechim — Kafka + "
    "Outbox Pattern — tafsilotlari bilan tushuntirilgan. Hujjat texnik developer va architect'lar "
    "uchun mo'ljallangan, lekin yuqori darajadagi qarorlar (qaysi yechim, qachon, nima xarajat) "
    "decision maker'lar uchun ham mavjud. "
)
p(
    "Asosiy yakuniy qaror: Hozirdan boshlab Kafka cluster + Outbox pattern o'rnatish, "
    "REST endpoint'lar 12+ oy parallel saqlanadi, 224 OTM o'z tezligida Kafka producer'ga "
    "o'tadi. Pilot bosqich olib tashlandi — loyiha hali production'ga ulanmagan, canary "
    "uchun real foydalanuvchi yo'q.",
    bold=True,
)

doc.add_page_break()

# ============================================================
# 1. KIRISH VA MAQSAD
# ============================================================
add_h("1. Kirish va maqsad", level=1)

add_h("1.1 Bu hujjat nima uchun yozildi", level=2)
p(
    "HEMIS-back (markaziy Spring Boot 4 modulli monolit) loyihasi 224 ta universitet "
    "(Univer — Yii2 PHP) bilan REST orqali ma'lumot almashadi. Hozirgi sinxron yondashuv "
    "ko'plab muammolarni keltirib chiqaradi: ma'lumot yo'qolishi, idempotency yo'qligi, "
    "observability past, back-channel mavjud emas, va h.k."
)
p(
    "Loyiha hali ishlab chiqarish (production) muhitiga ulanmagan — 224 OTM hozircha eski "
    "old-hemis CUBA Platform tizimi bilan ishlaydi. Bu greenfield holat — arxitekturani "
    "boshidanoq to'g'ri qurish imkonini beradi."
)
p(
    "Ushbu hujjat:"
)
bullets([
    "Hozirgi sync mexanizmini chuqur tushuntirib beradi",
    "8 ta aniq muammoni qanday yuzaga kelishini ko'rsatadi",
    "8 ta alternativ yechimni baholaydi (qaysi mezon bo'yicha qaysi yaxshi)",
    "Tanlangan yechim — Kafka + Outbox Pattern — tafsilotlari",
    "Univer va Hemis-back o'rtasidagi ma'lumot oqimini step-by-step ko'rsatadi",
    "Implementation reja, risk va konfiguratsiya namunalari",
])

add_h("1.2 Hujjat kim uchun", level=2)
table(
    ["O'quvchi roli", "Foyda"],
    [
        ["Backend developer (Hemis-back)", "Kod yozish uchun aniq pattern, joylashuv, kod misollar"],
        ["Frontend (PHP/Univer) developer", "Univer tomon nima o'zgarishi, qachon, qanday integratsiya"],
        ["Architect", "Qaror motivatsiyasi, alternativlar tahlili, trade-off matritsi"],
        ["Decision maker (PM, CTO)", "Vaqt, xarajat, risk hisoblash"],
        ["DevOps", "Docker, infra, monitoring, deploy zarurati"],
    ],
)

add_h("1.3 Yakuniy qaror (TL;DR)", level=2)
p(
    "Hozirdan boshlab Kafka + Outbox Pattern infrastruktura o'rnatiladi. Backend ichidagi "
    "har bir write operation atomic ravishda business jadval va outbox jadvalga yozadi. "
    "Background scheduler outbox jadvalni Kafka'ga yo'naltiradi. REST endpoint javobi "
    "o'zgarmaydi — 224 OTM PHP klient buzilmaydi.",
    bold=True,
)
p(
    "Pilot bosqich olib tashlandi — production'da real foydalanuvchi yo'q, canary o'rnida "
    "to'g'ri arxitektura bilan greenfield boshlash maqulroq.",
    italic=True,
)

doc.add_page_break()

# ============================================================
# 2. LOYIHA KONTEKSTI
# ============================================================
add_h("2. Loyiha konteksti", level=1)

add_h("2.1 Hemis-back nima qiladi", level=2)
p(
    "Hemis-back — O'zbekiston Oliy Ta'lim vazirligi tarkibidagi markaziy ma'lumotlar "
    "tizimi backend'i. Texnologik stack:"
)
bullets([
    "Java 25 LTS + Spring Boot 4.0.6 (modulli monolit)",
    "PostgreSQL 18 (master/replica) — asosiy ma'lumotlar bazasi",
    "Redis 7 — cache + session storage",
    "Liquibase 4.31 — schema migration",
    "Modul tuzilishi: common → domain → service → api-* → app",
    "Auxiliary DB: hemis_audit (alohida PostgreSQL — activity_log, error_log, login_log)",
])
p(
    "Hemis-back asosiy vazifalari: 224 OTM ma'lumotlarini markaziy ravishda saqlash, "
    "vazirlik va vazirlikka aloqador idoralar uchun hisobotlar tayyorlash, klassifikator "
    "va spravochnik ma'lumotlarni boshqarish, va tashqi tashkilotlar (BIMM, MyGov, GUVD, "
    "Soliq) bilan integratsiya."
)

add_h("2.2 Univer (224 OTM) nima qiladi", level=2)
p(
    "Univer — har bir universitet ichida deploy qilingan operativ tizim. Yii2 PHP "
    "freymvork'da yozilgan. 224 ta universitetda mustaqil ravishda ishlaydi — har OTM o'z "
    "PostgreSQL bazasi (hemis_337, hemis_401, hemis_NNN ...) bilan."
)
bullets([
    "Talaba ro'yxatga olish, kontrakt, akademik ma'lumotlar (baholar, GPA)",
    "O'qituvchilar ro'yxati, lavozim, akademik daraja",
    "Ilmiy maqolalar, loyihalar, dissertatsiyalar",
    "Diplom blanklari, sertifikatlar, hujjatlar",
    "Stipendiya, ish haqi, hisob-kitoblar (mahalliy)",
])
p(
    "Univer markaziy hisobot tayyorlamaydi — bu Hemis-back vazifasi. Univer faqat o'z OTM "
    "ichidagi operativ ish uchun. Markaziga ma'lumot REST orqali yuboriladi."
)

add_h("2.3 Ular qanday bog'langan", level=2)
code(
    "[224 ta OTM]                    [Hemis-back]\n"
    "hemis_337 (Yii2 PHP)  ──HTTPS──▶  /app/rest/v2/entities/hemishe_*\n"
    "hemis_401 (Yii2 PHP)  ──HTTPS──▶  /app/rest/v2/services/*\n"
    "hemis_NNN (Yii2 PHP)  ──HTTPS──▶  central PostgreSQL\n"
    "                                  (hemis_central + hemis_audit)\n\n"
    "Hemis-back ───────────❌──────────▶  Univer (back-channel YO'Q)",
    label="Hozirgi topologiya:",
)
p(
    "Bu yo'nalish — Univer → Hemis-back (1-yo'nalishli). Hemis-back o'zgarishlari "
    "(klassifikator yangilanish, status code qo'shilish) avtomatik ravishda 224 OTM'ga "
    "yetkazilmaydi. Har OTM o'zi cron orqali sport qilib tortib oladi."
)

add_h("2.4 Hozirgi miqdor (statistika)", level=2)
table(
    ["Mezon", "Qiymat"],
    [
        ["Universitetlar soni", "224 OTM"],
        ["Unique REST endpoint", "67 ta (33 entity + 32 service + 2 OAuth)"],
        ["Univer tomon caller class", "35 ta sync class (PHP)"],
        ["Hemis-back tomon controller", "123 ta (api-legacy modulda)"],
        ["O'rtacha sync hajmi (productionga ulanganda)", "~25,000 POST/day"],
        ["Peak time (semestr boshi)", "~50 OTM concurrent"],
        ["Auth grant", "OAuth password (deprecated, ADR-0005 da migration plan)"],
        ["Schema status", "FROZEN hemishe_* (refactor 224 OTM ta'sir qiladi)"],
    ],
)

doc.add_page_break()

# ============================================================
# 3. HOZIRGI SYNC ARXITEKTURASI
# ============================================================
add_h("3. Hozirgi sync arxitekturasi", level=1)

add_h("3.1 Univer tomon — qanday ma'lumot yuboradi", level=2)
p(
    "Univer tarkibida `common/components/hemis/sync/` papka 32 ta updater class'dan iborat. "
    "Har updater ma'lum entity uchun Hemis-back'ga POST qiladi. Markaziy klient — `HemisApi.php`."
)

p("Auth flow:", bold=True)
code(
    "// HemisApi.php:724-735\n"
    "$response = $this->_client->post('v2/oauth/token', [\n"
    "    'grant_type' => 'password',           // ← deprecated grant\n"
    "    'username'   => $otm_username,\n"
    "    'password'   => $otm_password,\n"
    "]);\n"
    "$accessToken = $response->getBody()->access_token;\n"
    "// Token har 12 soat o'rgatiladi (refresh)",
    label="PHP — Univer auth (HemisApi.php):",
)

p("Entity yuborish misoli:", bold=True)
code(
    "// EmployeeUpdater.php — soddalashtirilgan\n"
    "class EmployeeUpdater extends BaseApiUpdater {\n"
    "    public static function pushTeacher($teacher) {\n"
    "        $url = '/v2/entities/hemishe_ETeacher';\n"
    "        $data = [\n"
    "            'pinfl' => $teacher->pinfl,\n"
    "            'firstName' => $teacher->first_name,\n"
    "            'lastName' => $teacher->last_name,\n"
    "            'university' => ['id' => $teacher->university_id],\n"
    "            // ... 50+ field\n"
    "        ];\n"
    "        $response = self::client()->post($url, $data)->send();\n"
    "        // Network 503? Hemis-back deploy in progress?\n"
    "        // → silently lost (no retry, no DLQ, no outbox)\n"
    "    }\n"
    "}",
    label="PHP — Entity push (EmployeeUpdater.php):",
)

p(
    "Diqqat: bu fire-and-forget pattern. Network xatosi bo'lsa, ma'lumot yo'qoladi. PHP cron "
    "5 daqiqada bir takrorlasa ham, har retry'da yangi POST yuboriladi (idempotency yo'q — "
    "duplicate insert ehtimoli). Markaziy event log mavjud emas."
)

add_h("3.2 Hemis-back tomon — qanday qabul qiladi", level=2)
p("api-legacy modul 123 ta controller bilan barcha CUBA-style URL'larni qabul qiladi.")

code(
    "// api-legacy/.../TeacherEntityController.java\n"
    "@RestController\n"
    "@RequestMapping(\"/app/rest/v2/entities/hemishe_ETeacher\")\n"
    "public class TeacherEntityController {\n"
    "    private final TeacherService service;\n\n"
    "    @PostMapping\n"
    "    public ResponseEntity<?> create(@RequestBody Map<String, Object> body) {\n"
    "        // 1. CUBA convention bilan field'lar (LinkedHashMap)\n"
    "        // 2. Service'ga delegate\n"
    "        Teacher t = service.create(body);\n"
    "        // 3. CUBA-style javob (id, _entityName, _instanceName)\n"
    "        return ResponseEntity.ok(LegacyEntityAdapter.toMap(t));\n"
    "    }\n"
    "}",
    label="Java — Hemis-back receive (TeacherEntityController):",
)

code(
    "// service/.../TeacherService.java\n"
    "@Service\n"
    "@Transactional\n"
    "public class TeacherService {\n"
    "    public Teacher create(Map<String, Object> body) {\n"
    "        // 1. Validation (PINFL unique, etc.)\n"
    "        // 2. PostgreSQL INSERT\n"
    "        Teacher saved = repo.save(toEntity(body));\n"
    "        // 3. Return\n"
    "        return saved;\n"
    "        // ⚠ Faqat business INSERT — outbox/Kafka YO'Q (hozirgi holat)\n"
    "    }\n"
    "}",
    label="Java — Service layer (hozirgi):",
)

p(
    "Asosiy chiqib turgan kamchiliklar shu joyda: backend faqat o'z DB'ga yozadi. Audit, "
    "boshqa modul'lar, Univer back-channel — hammasi keyinchalik trigger orqali yoki cron "
    "orqali qilinadi (ishonchsiz)."
)

add_h("3.3 Real misol — Student create flow (step-by-step)", level=2)
code(
    "Step 1. Univer'da o'qituvchi yangi talaba qo'shadi:\n"
    "  hemis_337.student INSERT (Univer mahalliy DB)\n\n"
    "Step 2. Univer cron (5 min) StudentUpdater chaqiradi:\n"
    "  POST /app/rest/v2/entities/hemishe_EStudent\n"
    "  Authorization: Bearer eyJ...\n"
    "  Content-Type: application/json\n"
    "  Body: {pinfl: \"31234567890123\", firstName: \"Ali\", ...}\n\n"
    "Step 3. Hemis-back StudentEntityController qabul qiladi:\n"
    "  - Auth filter: JWT validate\n"
    "  - Controller: body → DTO\n"
    "  - Service: PINFL unique check + INSERT INTO hemishe_e_student\n"
    "  - Response: 200 OK + {id: 123, _entityName: \"hemishe_EStudent\", ...}\n\n"
    "Step 4. ❌ Audit log? — alohida @Async event listener (eventually)\n"
    "Step 5. ❌ Boshqa OTM uchun visible? — yo'q (har OTM mahalliy DB)\n"
    "Step 6. ❌ Read replica? — PostgreSQL physical replication (lag bor)",
    label="Hozirgi oqim (kamchilik bilan):",
)

add_h("3.4 Hozirgi 8 ta muammo (pain points)", level=2)
table(
    ["№", "Muammo", "Sabab", "Ta'sir"],
    [
        ["1", "Data loss", "Fire-and-forget POST, retry yo'q, DLQ yo'q", "Sync ishonchsiz, audit gap"],
        ["2", "Idempotency yo'q", "Idempotency-Key header yo'q, business_key upsert har controller'da o'zicha", "Network retry → duplicate insert"],
        ["3", "Concurrency (lost update)", "ETag/If-Match yo'q, JPA @Version ishlatilmaydi", "2 admin bir vaqt edit → biri yo'qoladi"],
        ["4", "Bulk inefficiency", "1000 student = 1000 POST = 1000 TCP handshake", "Peak time queue, thread starvation"],
        ["5", "Observability past", "Markaziy event log yo'q, debug 224 OTM SSH access", "Incident analysis qiyin"],
        ["6", "Back-channel YO'Q", "Hemis classifier yangilanish push qilinmaydi", "OTM cron pull = 1-24 soat lag"],
        ["7", "Schema drift", "Univer JSON yangi field silently dropped, contract test yo'q", "Ma'lumot yo'qolishi sezilmasdan"],
        ["8", "Schema coupling", "67 endpoint hemishe_* ga 1:1 mapped", "Refactor 224 OTM ta'sir riski"],
    ],
    col_widths=[Cm(1), Cm(3.5), Cm(5), Cm(5)],
)

doc.add_page_break()

# ============================================================
# 4. ALTERNATIVE YECHIMLAR
# ============================================================
add_h("4. Alternativ yechimlar (8 ta variant)", level=1)
p(
    "Quyida 8 ta yechim baholangan. Har biri uchun: nima qiladi, afzalliklari, kamchiliklari, "
    "qaror (qabul qilingan / qisman / rad etilgan)."
)

# Alt 1
add_h("4.1 Status quo — REST only (DOING NOTHING)", level=2)
p("Tasvir: Hech nima o'zgartirmaymiz. REST POST davom etadi.")
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["O'zgartirish kerak emas", "Data loss eksponensial o'sadi"],
        ["Hozirgi 224 OTM kod ishlaydi", "Bulk POST 50 OTM peak'da queue"],
        ["", "Observability past — incident debug og'ir"],
        ["", "Back-channel yo'qligi: yangi classifier 1-24 soat lag bilan"],
        ["", "Audit gap — ma'lumot yo'qolishi sezilmasdan o'tadi"],
    ],
)
p("Qaror: REJECTED — scale uchun yetarli emas.", bold=True)

# Alt 2
add_h("4.2 REST hardening (Idempotency-Key, ETag, bulk, webhook)", level=2)
p(
    "Tasvir: REST endpoint'larga Idempotency-Key header (Stripe pattern), ETag/If-Match "
    "(RFC 7232), bulk endpoint, webhook back-channel qo'shamiz."
)
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["REST stack saqlanadi", "Webhook delivery zaif (network drop = lost)"],
        ["Standard pattern (Stripe, GitHub)", "Replay imkoni cheklangan"],
        ["Klient tomon kichik o'zgartirish", "Order kafolati yo'q (concurrent webhook)"],
        ["", "Kafka kelganda hammasi keraksiz qatlam bo'ladi"],
        ["", "Outbox pattern oldi olinmaydi (data loss qoladi)"],
    ],
)
p("Qaror: PARTIAL/REJECTED — Kafka + Outbox kelganda bu qatlamlar keraksiz. ", bold=True)
p(
    "Idempotency-Key REST klient tomonda foydali bo'lishi mumkin (224 OTM hali REST yuborayotganda), "
    "lekin Kafka exactly-once consumer qoplay oladi.",
    italic=True,
)

# Alt 3
add_h("4.3 Big-bang Kafka migration (REST drop boshidanoq)", level=2)
p("Tasvir: REST endpoint'larni o'chirib, faqat Kafka producer/consumer qoldiramiz.")
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["Toza arxitektura", "224 OTM PHP team bir vaqtda yangilanishi kerak"],
        ["Bitta protocol", "Backward compat kafolati buziladi"],
        ["Kafka native semantics", "Rollback yo'q (REST allaqachon o'chirilgan)"],
        ["", "Kafka cluster ops debt katta (zero-experience)"],
    ],
)
p("Qaror: REJECTED — REST 12+ oy saqlanishi mantiqiy zarurat (224 OTM tezligi turlicha).", bold=True)

# Alt 4
add_h("4.4 Database logical replication (PostgreSQL native)", level=2)
p("Tasvir: Univer DB → Hemis-back DB to'g'ridan-to'g'ri replikatsiya (pg_logical).")
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["Application kod o'zgarmaydi", "224 ta DB ↔ 1 markaz schema mismatch"],
        ["DB-level kafolat", "FK noaniqlik (Univer.faculty_id ↔ markaz.faculty.code)"],
        ["", "DDL change'larni 224 ta DB'ga sinxron qo'llash"],
        ["", "pg_logical replication slot 224 ta — master CPU/memory load"],
    ],
)
p("Qaror: REJECTED — schema heterogeneity DB-level repl'ni qo'llab-quvvatlamaydi.", bold=True)

# Alt 5
add_h("4.5 GraphQL Federation", level=2)
p("Tasvir: Apollo Federation, schema-as-contract.")
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["Klient flexibility", "Sync write-heavy (read-heavy emas) — afzallik yo'qoladi"],
        ["Schema tip-safety", "N+1 risk (federation resolvers chain)"],
        ["", "PHP GraphQL klient zaif ekosistem"],
    ],
)
p("Qaror: REJECTED — wrong tool for write-heavy use case.", bold=True)

# Alt 6
add_h("4.6 gRPC streaming", level=2)
p("Tasvir: HTTP/2 multiplexing, binary protobuf.")
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["Binary efficient", "PHP gRPC ekosistemi zaif (rdkafka extension talab)"],
        ["HTTP/2 multiplexing", "Schema evolution Kafka'dan zaif"],
    ],
)
p("Qaror: PARTIAL — Kafka producer'larda binary serialization (Avro) ishlatamiz.", bold=True)

# Alt 7
add_h("4.7 Webhook-only (push-pull)", level=2)
p("Tasvir: Hemis → Univer push, Univer → Hemis pull. Sodda HTTP.")
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["Sodda", "Webhook delivery guarantee zaif"],
        ["", "Replay yo'q, history saqlanmaydi"],
        ["", "Order kafolati yo'q (concurrent webhook)"],
    ],
)
p("Qaror: REJECTED — Kafka topic to'liq qoplaydi (durable, replayable, ordered).", bold=True)

# Alt 8 — TANLANGAN
add_h("4.8 Kafka + Outbox Pattern ★ (TANLANGAN)", level=2)
p("Tasvir: ", bold=True)
p(
    "Backend ichida har bir write operation atomic ravishda business jadval va outbox jadvalga "
    "yozadi. Background scheduler outbox jadvaldan o'qib Kafka topic'larga yo'naltiradi. "
    "REST endpoint'lar saqlanadi (parallel). Schema Registry (Apicurio) JSON Schema validation. "
    "Pilot bosqich olib tashlandi (greenfield)."
)
table(
    ["Afzalligi", "Kamchiligi"],
    [
        ["Data durability — at-least-once + idempotent consumer = effectively exactly-once", "4-6 hafta investment Kafka foundation"],
        ["Observability — Kafka topic = central event log (7+ day retention)", "Yangi konsept team uchun (outbox pattern)"],
        ["Decoupling — Univer va Hemis-back deploy mustaqil", "Schema evolution governance kerak"],
        ["Audit trail — har event sourceable", "Lokal docker overhead (~512MB RAM)"],
        ["Bulk efficient — Kafka batch produce (1000 events ~100ms)", ""],
        ["Back-channel — Hemis → Univer Kafka topic (lag < 5s)", ""],
        ["Replay — issue paydo bo'lsa 24h-7d replay", ""],
        ["Strangler-friendly — REST + Kafka parallel", ""],
    ],
)
p("Qaror: ACCEPTED ★", bold=True)

add_h("4.9 Comparison matrix (mezon bo'yicha)", level=2)
table(
    [
        "Mezon",
        "Status quo",
        "REST hardening",
        "Big-bang Kafka",
        "DB replication",
        "GraphQL",
        "gRPC",
        "Webhook",
        "Kafka+Outbox ★",
    ],
    [
        ["Data durability", "✗", "△", "✓", "✓", "△", "△", "✗", "✓"],
        ["Backward compat", "✓", "✓", "✗", "✓", "△", "△", "✓", "✓"],
        ["Observability", "✗", "△", "✓", "△", "△", "△", "△", "✓"],
        ["Bulk efficiency", "✗", "△", "✓", "✓", "✗", "△", "✗", "✓"],
        ["Back-channel", "✗", "△", "✓", "✗", "△", "△", "△", "✓"],
        ["Replay imkoni", "✗", "✗", "✓", "△", "✗", "✗", "✗", "✓"],
        ["Schema governance", "✗", "△", "✓", "✗", "✓", "△", "✗", "✓"],
        ["Implementation cost", "0", "kichik", "katta", "katta", "katta", "katta", "kichik", "o'rta"],
        ["Risk darajasi", "high", "low", "very high", "high", "high", "med", "med", "low"],
    ],
)
p("Belgilar: ✓ = qoplaydi, △ = qisman, ✗ = qoplamaydi", italic=True, size=9)

doc.add_page_break()

# ============================================================
# 5. TANLANGAN YECHIM — KAFKA + OUTBOX
# ============================================================
add_h("5. Tanlangan yechim — Kafka + Outbox Pattern", level=1)

add_h("5.1 Outbox Pattern nima", level=2)
p(
    "Outbox Pattern (Chris Richardson, \"Microservices Patterns\", 2018) — distributed system'da "
    "atomic dual-write muammosini hal qiluvchi pattern. Asosiy g'oya:"
)
numbered([
    "Service ma'lumotni DB'ga yozayotganda, BIRGA outbox jadvalga ham event yozadi (bitta SQL transaction'da).",
    "Background process outbox jadvalni o'qib Kafka (yoki boshqa message broker)'ga yo'naltiradi.",
    "Outbox jadvalda muvaffaqiyatli yuborilgan event'lar published_at = NOW() bilan belgilanadi.",
    "Bu — \"transactional outbox\" — ma'lumot yo'qolmaydi: DB commit bo'lsa Kafka'ga yetkazilishi kafolatlanadi (eventually).",
])

p("Klassik dual-write antipattern (XATO):", bold=True)
code(
    "// ✗ XATO — dual-write race condition\n"
    "@Transactional\n"
    "public Student create(StudentDto dto) {\n"
    "    Student s = repo.save(toEntity(dto));      // DB INSERT\n"
    "    kafkaTemplate.send(\"student-events\", s);  // Kafka send\n"
    "    return s;\n"
    "}\n"
    "// Muammo: Kafka send xato bo'lsa-da, DB commit bo'ladi → inconsistency\n"
    "// Yoki: DB commit xato bo'lsa-da, Kafka send oldin bo'ldi → ghost event",
    label="Antipattern (xato yo'l):",
)

p("Outbox pattern (TO'G'RI):", bold=True)
code(
    "// ✓ TO'G'RI — outbox pattern\n"
    "@Transactional\n"
    "public Student create(StudentDto dto) {\n"
    "    Student s = repo.save(toEntity(dto));      // INSERT 1\n"
    "    outboxPublisher.publish(                    // INSERT 2 (same TX!)\n"
    "        \"student\", s.getId().toString(),\n"
    "        \"created\", s);\n"
    "    return s;\n"
    "    // Bitta transaction'da: ikkala INSERT atomic\n"
    "}\n\n"
    "// Background:\n"
    "@Scheduled(fixedDelay = 5000)\n"
    "@Transactional\n"
    "public void publishPending() {\n"
    "    List<OutboxEvent> pending = repo.findUnpublished(100);\n"
    "    for (OutboxEvent e : pending) {\n"
    "        kafkaTemplate.send(topic, key, payload).get();\n"
    "        e.markPublished();  // UPDATE outbox SET published_at = NOW()\n"
    "    }\n"
    "}",
    label="Outbox pattern (to'g'ri yo'l):",
)

add_h("5.2 Kafka roli", level=2)
bullets([
    "Distributed log — har event ketma-ket saqlanadi (offset bilan)",
    "Durable buffer — default 7 kun retention (configurable)",
    "Multiple consumer support — bir topic'ni ko'p consumer mustaqil o'qiy oladi",
    "Replay capability — eski event'larni qaytadan o'qish (debugging, recovery)",
    "Partition orqali ordering — bir kalit (aggregate_id) doim bir partitionda → kafolatli order",
    "Compaction — classifier topic'lari uchun (faqat oxirgi qiymat saqlanadi)",
])

add_h("5.3 Schema Registry (Apicurio)", level=2)
p(
    "Apicurio Registry — Red Hat tomonidan ishlab chiqilgan ochiq manbali schema management "
    "tizimi. Kafka topic'lariga JSON Schema yoki Avro schema bog'laydi. Foyda:"
)
bullets([
    "Backward compatibility check — yangi schema eski klient'larni buzmaydi",
    "Forward compatibility check — eski schema yangi klient'lar bilan ishlaydi",
    "Schema versioning — har breaking change yangi versiya",
    "CI gate — schema PR avtomatik tekshiriladi",
    "Open source (Apache 2.0), Confluent Schema Registry'dan litsenziya jihatdan free",
])

add_h("5.4 KRaft mode (no Zookeeper)", level=2)
p(
    "Apache Kafka 3.5+ KRaft (Kafka Raft) consensus protocol bilan ishlaydi. Zookeeper "
    "Apache 3.8'dan deprecated, 4.0'da olib tashlanadi. Afzalliklari:"
)
bullets([
    "50% kam resource (Zookeeper alohida JVM emas)",
    "Sodda deploy (1 ta service, 2 ta emas)",
    "Tezroq partition leader election",
    "Production scale: 1M+ partitions support",
])

doc.add_page_break()

# ============================================================
# 6. KOMPONENTLAR VA VAZIFALAR
# ============================================================
add_h("6. Komponentlar va vazifalar (detailed)", level=1)

add_h("6.1 Apache Kafka cluster", level=2)
p("Lokal dev: 1 ta broker (single-node KRaft)")
p("Production: 3 ta broker (HA, partition replication factor 3)")
table(
    ["Param", "Lokal", "Production"],
    [
        ["Broker soni", "1", "3"],
        ["Replication factor", "1", "3"],
        ["Min ISR", "1", "2"],
        ["acks (producer)", "all", "all"],
        ["Retention", "7 day", "7 day (events), forever (compacted)"],
        ["Disk", "Docker volume", "Separate SSD per broker"],
    ],
)

add_h("6.2 Apicurio Schema Registry", level=2)
p("Lokal: in-memory (apicurio-registry-mem). Production: PostgreSQL backed.")
p("REST API: /apis/registry/v2/groups/{group}/artifacts/{artifactId}")

add_h("6.3 Outbox jadval (V015 Liquibase migration)", level=2)
code(
    "CREATE TABLE outbox_event (\n"
    "    id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
    "    aggregate_type  VARCHAR(50) NOT NULL,    -- 'student', 'teacher', 'classifier'\n"
    "    aggregate_id    VARCHAR(100) NOT NULL,   -- entity ID (UUID, Long, code)\n"
    "    event_type      VARCHAR(50) NOT NULL,    -- 'created', 'updated', 'deleted'\n"
    "    payload         JSONB NOT NULL,          -- entity snapshot\n"
    "    schema_version  INT NOT NULL DEFAULT 1,\n"
    "    occurred_at     TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,\n"
    "    published_at    TIMESTAMP,               -- NULL = pending\n"
    "    retry_count     INT NOT NULL DEFAULT 0,\n"
    "    last_error      TEXT,\n"
    "    correlation_id  VARCHAR(100),            -- distributed tracing\n"
    "    causation_id    VARCHAR(100)             -- root cause event\n"
    ");\n\n"
    "CREATE INDEX idx_outbox_unpublished ON outbox_event (occurred_at)\n"
    "    WHERE published_at IS NULL;\n"
    "CREATE INDEX idx_outbox_aggregate ON outbox_event (aggregate_type, aggregate_id, occurred_at);",
    label="V015_create_outbox_event.sql:",
)

add_h("6.4 OutboxEventPublisher (helper)", level=2)
p("Service layer'dan chaqiriladi. Domain transaction ichida ishlaydi (atomic write).")
code(
    "@Component\n"
    "@RequiredArgsConstructor\n"
    "public class OutboxEventPublisher {\n"
    "    private final OutboxEventRepository repo;\n"
    "    private final ObjectMapper objectMapper;\n\n"
    "    public void publish(String aggregateType, String aggregateId,\n"
    "                        String eventType, Object payload) {\n"
    "        try {\n"
    "            String json = objectMapper.writeValueAsString(payload);\n"
    "            OutboxEvent e = OutboxEvent.builder()\n"
    "                .aggregateType(aggregateType)\n"
    "                .aggregateId(aggregateId)\n"
    "                .eventType(eventType)\n"
    "                .payload(json)\n"
    "                .schemaVersion(1)\n"
    "                .build();\n"
    "            repo.save(e);\n"
    "        } catch (JsonProcessingException ex) {\n"
    "            throw new IllegalStateException(\"Failed to serialize\", ex);\n"
    "        }\n"
    "    }\n"
    "}",
    label="OutboxEventPublisher.java:",
)

add_h("6.5 OutboxKafkaPublisher (scheduled)", level=2)
p("Background poll. Har 5 sekundda outbox jadvalni tekshiradi, Kafka'ga yo'naltiradi.")
code(
    "@Component\n"
    "@RequiredArgsConstructor\n"
    "@Slf4j\n"
    "public class OutboxKafkaPublisher {\n"
    "    private final OutboxEventRepository repo;\n"
    "    private final KafkaTemplate<String, String> kafkaTemplate;\n\n"
    "    @Scheduled(fixedDelayString = \"${hemis.outbox.poll-interval-ms:5000}\")\n"
    "    @Transactional\n"
    "    public void publishPending() {\n"
    "        List<OutboxEvent> pending = repo.findUnpublished(\n"
    "            PageRequest.of(0, 100));\n"
    "        for (OutboxEvent e : pending) {\n"
    "            try {\n"
    "                String topic = \"hemis.\" + e.getAggregateType() +\n"
    "                               \".events.v\" + e.getSchemaVersion();\n"
    "                kafkaTemplate.send(topic,\n"
    "                                   e.getAggregateId(),\n"
    "                                   e.getPayload())\n"
    "                              .get(5, TimeUnit.SECONDS);\n"
    "                e.setPublishedAt(LocalDateTime.now());\n"
    "            } catch (Exception ex) {\n"
    "                e.setRetryCount(e.getRetryCount() + 1);\n"
    "                e.setLastError(ex.getMessage());\n"
    "                if (e.getRetryCount() > 5) {\n"
    "                    log.error(\"Outbox {} → DLQ\", e.getId(), ex);\n"
    "                    kafkaTemplate.send(\"hemis.dlq.v1\",\n"
    "                                       e.getAggregateId(), e.getPayload());\n"
    "                    e.setPublishedAt(LocalDateTime.now());\n"
    "                }\n"
    "            }\n"
    "        }\n"
    "    }\n"
    "}",
    label="OutboxKafkaPublisher.java:",
)

doc.add_page_break()

# ============================================================
# 7. TOPIC STRUKTURA
# ============================================================
add_h("7. Topic struktura", level=1)

add_h("7.1 Naming konvensiya", level=2)
code(
    "Pattern: hemis.{aggregate_type}.events.v{schema_version}\n\n"
    "Hemis-back tomon (Stage 1):\n"
    "  hemis.student.events.v1\n"
    "  hemis.teacher.events.v1\n"
    "  hemis.classifier.h_position.events.v1\n"
    "  hemis.classifier.h_construction_material.events.v1\n"
    "  hemis.publication.events.v1\n"
    "  hemis.project.events.v1\n"
    "  hemis.audit.events.v1\n\n"
    "Univer tomondan (Stage 2):\n"
    "  univer.student.events.v1\n"
    "  univer.teacher.events.v1\n"
    "  univer.publication.events.v1\n\n"
    "DLQ (Dead Letter Queue):\n"
    "  hemis.dlq.v1",
    label="Topic naming:",
)

add_h("7.2 Partition strategy", level=2)
table(
    ["Topic", "Partition kalit", "Sabab"],
    [
        ["hemis.student.events.v1", "aggregate_id (student UUID)", "Bir student'ning event'lari ketma-ket kelishi"],
        ["univer.student.events.v1", "university_code", "Per-OTM ordering"],
        ["hemis.classifier.*.v1", "code", "Klassifikator kalit kichik domain"],
        ["hemis.audit.events.v1", "aggregate_type", "Audit grouped by domain"],
        ["hemis.dlq.v1", "aggregate_id", "Easy investigation"],
    ],
)
p("Partition soni: 3-6 (lokal), 12-24 (production) — har topic uchun.")

add_h("7.3 Retention policy", level=2)
table(
    ["Topic turi", "Retention", "Sabab"],
    [
        ["Domain events (student, teacher, ...)", "7 day", "Replay imkoni, debug"],
        ["Audit events", "30 day", "Compliance, audit trail"],
        ["Classifier events", "Compacted (forever)", "Faqat oxirgi qiymat — full snapshot"],
        ["DLQ", "30 day", "Manual reprocess imkoni"],
    ],
)

add_h("7.4 Schema versioning strategy", level=2)
p(
    "Backward compatibility schema breaking change'larda yangi `.v2` topic. Eski `.v1` topic "
    "konsumerlar tomonidan o'qilishi davom etadi (ma'lum vaqt). Bu pattern Confluent va Stripe "
    "tomonidan tavsiya etiladi."
)
code(
    "Schema evolution misoli:\n\n"
    "v1 schema:\n"
    "  {pinfl, firstName, lastName}\n\n"
    "v2 schema (backward-compatible — yangi optional field):\n"
    "  {pinfl, firstName, lastName, middleName?}  -- Apicurio: BACKWARD\n\n"
    "v2 schema (BREAKING — required field):\n"
    "  {pinfl, firstName, lastName, middleName}   -- yangi topic kerak\n"
    "  → hemis.student.events.v2\n"
    "  → producer dual-publish (.v1 + .v2) ma'lum vaqt\n"
    "  → consumer .v1 → .v2 migrate o'z tezligida\n"
    "  → .v1 deprecate (3-6 oy keyin)",
    label="Schema versioning:",
)

doc.add_page_break()

# ============================================================
# 8. UNIVER ↔ HEMIS-BACK DATA EXCHANGE
# ============================================================
add_h("8. Univer ↔ Hemis-back data exchange (full picture)", level=1)

add_h("8.1 Stage 1 (hozir, 4-6 hafta)", level=2)
p(
    "REST endpoint'lar saqlanadi. Backend internal'da outbox + Kafka qo'shiladi. 224 OTM "
    "hech narsani sezmaydi — REST javob shape o'zgarmaydi."
)
code(
    "224 OTM (Univer Yii2 PHP)\n"
    "    │\n"
    "    │ 1. POST /v2/entities/hemishe_EStudent (REST, JWT)\n"
    "    ▼\n"
    "Hemis-back (Spring Boot)\n"
    "    │\n"
    "    │ 2. StudentEntityController @PostMapping\n"
    "    │ 3. StudentService @Transactional {\n"
    "    │      repo.save(student);          ← DB INSERT 1\n"
    "    │      outbox.publish(\"student\",     ← DB INSERT 2 (atomic)\n"
    "    │                     id, \"created\", student);\n"
    "    │    }\n"
    "    │ 4. Response 200 OK + CUBA-style JSON\n"
    "    │\n"
    "    │ 5. Background: OutboxKafkaPublisher (every 5s)\n"
    "    │      kafkaTemplate.send(\"hemis.student.events.v1\", id, payload);\n"
    "    │\n"
    "    ▼\n"
    "Kafka cluster\n"
    "    │\n"
    "    │ 6. audit-event-consumer\n"
    "    │      → hemis_audit DB INSERT INTO activity_log\n"
    "    │\n"
    "    │ 7. (kelajakda) read-replica-projector\n"
    "    │      → Cassandra/Elasticsearch projection\n"
    "    │\n"
    "    │ 8. (Stage 2) Univer-subscriber-consumer\n"
    "    │      → Univer Webhook back-channel",
    label="Stage 1 oqim:",
)

add_h("8.2 Stage 2 (kelajak, OTM ready)", level=2)
p(
    "Univer tomon Kafka producer integratsiyasi. PHP rdkafka extension orqali. Per-OTM "
    "feature flag — har OTM o'z tezligida o'tadi. REST parallel ishlaydi (rollback xavfsizligi)."
)
code(
    "// Univer common/components/hemis/HemisApi.php\n"
    "public function syncStudent($student) {\n"
    "    if (Config::get(Config::CONFIG_USE_KAFKA)) {\n"
    "        // Yangi Kafka rejimi\n"
    "        return $this->kafkaPublisher->publish(\n"
    "            'univer.student.events.v1',\n"
    "            $student->pinfl,\n"
    "            $student->toArray()\n"
    "        );\n"
    "    }\n"
    "    // Eski rejim — REST POST (default, 12+ oy parallel)\n"
    "    return $this->_client\n"
    "        ->post('/v2/entities/hemishe_EStudent', $student->toArray())\n"
    "        ->send();\n"
    "}",
    label="PHP — Stage 2 producer (Univer tomon):",
)

p("Univer tomonda ham outbox jadval qo'shiladi:")
code(
    "// Univer hemis_NNN.sql\n"
    "CREATE TABLE outbox_event (\n"
    "    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n"
    "    aggregate_type VARCHAR(50),\n"
    "    payload JSONB,\n"
    "    occurred_at TIMESTAMP DEFAULT NOW(),\n"
    "    published_at TIMESTAMP\n"
    ");\n\n"
    "// PHP service\n"
    "DB::transaction(function () use ($student) {\n"
    "    Student::create($student);            // DB INSERT 1\n"
    "    OutboxEvent::create([                 // DB INSERT 2 (same TX)\n"
    "        'aggregate_type' => 'student',\n"
    "        'payload' => json_encode($student),\n"
    "    ]);\n"
    "});",
    label="PHP — Univer outbox:",
)

add_h("8.3 Back-channel (Hemis → Univer)", level=2)
p(
    "Hemis-back klassifikator yangilanish (yangi `h_position` qiymati) — avtomatik 224 OTM'ga "
    "Kafka topic orqali yetkaziladi. Lag < 5 sekund."
)
code(
    "Hemis admin panel → 'h_position' yangi qiymat qo'shadi:\n"
    "    │\n"
    "    ▼\n"
    "ClassifierService.create('h_position', 'METHOD', 'Metodist'):\n"
    "    │  @Transactional\n"
    "    │  repo.save(...);                                    INSERT 1\n"
    "    │  outbox.publish(\"classifier.h_position\",            INSERT 2\n"
    "    │                  \"METHOD\", \"created\", classifier);\n"
    "    ▼\n"
    "OutboxKafkaPublisher (5s):\n"
    "    │  topic = \"hemis.classifier.h_position.events.v1\"\n"
    "    │  kafka.send(topic, \"METHOD\", payload);\n"
    "    ▼\n"
    "Kafka cluster (compacted topic)\n"
    "    │\n"
    "    ▼\n"
    "Univer subscriber consumer (har OTM):\n"
    "    │  consumer.subscribe([\"hemis.classifier.h_position.events.v1\"]);\n"
    "    │  Har event uchun: PHP DB UPSERT INTO h_position\n"
    "    │  → Univer mahalliy DB'ga yangi klassifikator yetib boradi\n"
    "    ▼\n"
    "OTM admin/talaba endi yangi 'METHOD' qiymatini ko'radi\n"
    "(eski cron pull o'rnida real-time push)",
    label="Back-channel oqim:",
)

add_h("8.4 Comparison: REST vs Kafka", level=2)
table(
    ["Mezon", "REST (hozirgi)", "Kafka (yangi)"],
    [
        ["Sync mode", "Synchronous (blocking)", "Asynchronous (non-blocking)"],
        ["Delivery guarantee", "Best-effort (HTTP retry klient'da)", "At-least-once (Kafka offset)"],
        ["Latency p99", "~500ms", "~50ms (Kafka producer)"],
        ["Bulk efficiency", "1 row = 1 HTTP", "Batch (1000 rows ~100ms)"],
        ["Order kafolati", "Yo'q (concurrent POST)", "Per-partition order kafolati"],
        ["Replay", "Mumkin emas", "Possible (offset reset)"],
        ["Back-channel", "Webhook (zaif)", "Topic broadcast"],
        ["Schema evolution", "Manual (backward break risk)", "Schema Registry (BACKWARD policy)"],
        ["Observability", "Per-OTM SSH log", "Markaziy Kafka topic + Grafana"],
        ["Audit trail", "Eventual @Async", "Inline (atomic)"],
    ],
)

doc.add_page_break()

# ============================================================
# 9. REAL MISOL — STUDENT CREATE FLOW
# ============================================================
add_h("9. Real misol — Student create flow", level=1)

add_h("9.1 Stage 1 (hozir, REST + outbox)", level=2)
code(
    "T+0ms   [Univer hemis_337]\n"
    "        Talaba ro'yxatga olinadi (mahalliy DB):\n"
    "        INSERT INTO hemis_337.student (...);\n\n"
    "T+5min  [Univer cron]\n"
    "        StudentUpdater::pushAllUnsynced():\n"
    "          POST https://hemis.uz/v2/entities/hemishe_EStudent\n"
    "          Authorization: Bearer eyJhbGc...\n"
    "          Body: {pinfl, firstName, lastName, university: {id: 337}, ...}\n\n"
    "T+5min  [Hemis-back StudentEntityController]\n"
    "        @PostMapping receive request\n"
    "        body → StudentDto\n"
    "        studentService.create(dto):\n\n"
    "          @Transactional begin\n"
    "          ─── INSERT INTO hemishe_e_student (id=42, pinfl, ...);\n"
    "          ─── INSERT INTO outbox_event (\n"
    "                  aggregate_type='student',\n"
    "                  aggregate_id='42',\n"
    "                  event_type='created',\n"
    "                  payload='{\"id\":42, \"pinfl\":...}',\n"
    "                  schema_version=1,\n"
    "                  occurred_at=NOW(),\n"
    "                  published_at=NULL\n"
    "              );\n"
    "          @Transactional commit  ← ikkala INSERT atomic\n\n"
    "        Response 200 OK\n"
    "          {id: 42, _entityName: 'hemishe_EStudent', _instanceName: '...'}\n\n"
    "T+5min  [Univer]\n"
    "        Response qabul qiladi, mahalliy student.synced = true\n\n"
    "═══════════════════════════════════════════════════════════════════\n"
    "Background (alohida thread, REST request'dan keyin):\n"
    "═══════════════════════════════════════════════════════════════════\n\n"
    "T+5min+5s  [Hemis-back OutboxKafkaPublisher]\n"
    "           @Scheduled(fixedDelay=5000):\n"
    "           SELECT * FROM outbox_event WHERE published_at IS NULL\n"
    "                  ORDER BY occurred_at LIMIT 100;\n"
    "           → 1 ta event topildi (id=42)\n\n"
    "           kafkaTemplate.send(\n"
    "               topic='hemis.student.events.v1',\n"
    "               key='42',\n"
    "               value='{\"id\":42, \"pinfl\":..., \"eventType\":\"created\"}'\n"
    "           ).get();\n\n"
    "           UPDATE outbox_event SET published_at = NOW() WHERE id=...;\n\n"
    "T+5min+10s [Audit consumer]\n"
    "           @KafkaListener(topics='hemis.student.events.v1')\n"
    "           public void onStudentEvent(StudentEvent e):\n"
    "             auditDb.insert(activity_log, ...);\n\n"
    "T+5min+10s [Read replica projector — Stage 3]\n"
    "           @KafkaListener:\n"
    "             elasticsearch.index('students', e.id, e);\n\n"
    "T+5min+10s [Univer subscriber — Stage 2]\n"
    "           Boshqa universitet universitet_code=337 student'ni ko'rmasligi\n"
    "           kerak. Faqat 337 OTM o'zining boshqa instance'lari (HA) uchun.",
    label="Step-by-step (Stage 1):",
)

add_h("9.2 Stage 2 (kelajak, Univer Kafka producer)", level=2)
code(
    "T+0ms   [Univer hemis_337 Yii2]\n"
    "        Talaba ro'yxatga olinadi (mahalliy DB):\n"
    "        $student = Student::create($data);\n"
    "        // Outbox jadvalga ham yoziladi (atomic transaction):\n"
    "        OutboxEvent::create([\n"
    "            'aggregate_type' => 'student',\n"
    "            'aggregate_id' => $student->id,\n"
    "            'event_type' => 'created',\n"
    "            'payload' => json_encode($student),\n"
    "        ]);\n\n"
    "T+5s    [Univer KafkaPublisher cron]\n"
    "        SELECT * FROM outbox_event WHERE published_at IS NULL;\n"
    "        rdkafka_publish('univer.student.events.v1',\n"
    "                        '337-' . $student->id,\n"
    "                        json_encode($student));\n\n"
    "T+5s    [Hemis-back consumer]\n"
    "        @KafkaListener(topics='univer.student.events.v1')\n"
    "        public void consume(StudentEvent e):\n"
    "          studentService.upsertFromUniver(e);\n"
    "          // INSERT INTO hemishe_e_student (...)\n"
    "          // ON CONFLICT (pinfl) DO UPDATE SET ...\n"
    "          // ALSO: outbox.publish (re-broadcast hemis topic)\n"
    "        // 224 OTM bilan o'zaro audit log markazda\n"
    "        // No more REST POST!",
    label="Step-by-step (Stage 2):",
)

doc.add_page_break()

# ============================================================
# 10. IMPLEMENTATION REJA
# ============================================================
add_h("10. Implementation reja (Stage 1 — 4-6 hafta)", level=1)

add_h("10.1 10 ta atomik qadam", level=2)
table(
    ["№", "Modul", "Fayl", "Maqsad", "Vaqt"],
    [
        ["1", "root", "docker-compose.yml", "Kafka + Apicurio + Kafka-UI servisi", "30 min"],
        ["2", "service", "build.gradle.kts", "spring-kafka dep", "10 min"],
        ["3", "domain", "V015_create_outbox_event.sql + rollback", "Outbox jadval", "20 min"],
        ["4", "domain", "db.changelog-master.yaml", "V015 entry", "5 min"],
        ["5", "domain", "OutboxEvent.java", "JPA entity", "15 min"],
        ["6", "domain", "OutboxEventRepository.java", "Spring Data", "10 min"],
        ["7", "service", "OutboxEventPublisher.java", "Helper (atomic write)", "20 min"],
        ["8", "service", "OutboxKafkaPublisher.java", "Scheduled poll → Kafka", "30 min"],
        ["9", "app", "application.yml", "Kafka config", "10 min"],
        ["10", "app", "HemisApplication.java", "@EnableScheduling", "5 min"],
    ],
    col_widths=[Cm(0.8), Cm(1.5), Cm(5.5), Cm(5), Cm(1.5)],
)
p("Jami: ~3 soat code, +1 soat build/test verify, +2 soat docs = ~6 soat.")

add_h("10.2 Build verification checkpoints", level=2)
numbered([
    "Step 1 (docker-compose) — `docker compose config` syntax check",
    "Step 2-4 (gradle + migration + master) — `./gradlew :domain:liquibaseValidate :service:compileJava`",
    "Step 5-7 (entity + repo + helper) — `./gradlew :service:compileJava`",
    "Step 8 (publisher) — `./gradlew :service:compileJava` + unit test",
    "Step 9-10 (config + bootstrap) — `./gradlew :app:bootRun --dry-run`",
])

add_h("10.3 Birinchi domain integratsiyasi (Step 11+, alohida)", level=2)
p(
    "Yuqoridagi 10 ta qadam — faqat infrastruktura. Birinchi haqiqiy event publishing alohida "
    "Step 11+ orqali kiradi. Tavsiya: sodda klassifikator (`h_position`) bilan boshlash:"
)
bullets([
    "Step 11: ClassifierService.create() — outbox.publish() chaqiriq qo'shish",
    "Step 12: Topic create — `hemis.classifier.h_position.events.v1`",
    "Step 13: Test — POST classifier → outbox INSERT → Kafka → consumer log",
    "Step 14: Audit consumer (hemis_audit DB) ulash",
    "Step 15: 2-domain (StudentService) integratsiya",
])

add_h("10.4 Test strategiya", level=2)
table(
    ["Test darajasi", "Stack", "Maqsad"],
    [
        ["Unit", "JUnit + Mockito", "OutboxEventPublisher logic"],
        ["Unit", "JUnit + Mockito", "OutboxKafkaPublisher retry/DLQ logic"],
        ["Integration", "Testcontainers + Kafka + PostgreSQL", "End-to-end outbox → Kafka send"],
        ["Integration", "Testcontainers", "Schema validation (Apicurio)"],
        ["Manual", "docker compose + curl", "REST POST → Kafka topic verification"],
    ],
)

doc.add_page_break()

# ============================================================
# 11. RISK VA MITIGATION
# ============================================================
add_h("11. Risk va mitigation", level=1)

table(
    ["Risk", "Ehtimol", "Ta'sir", "Mitigation"],
    [
        [
            "Outbox queue ortib ketadi (Kafka downtime)",
            "Past",
            "Yuqori (PostgreSQL disk yetmaydi)",
            "Retry policy + DLQ topic. retry_count > 5 → DLQ + alert. Disk monitoring.",
        ],
        [
            "Schema breaking change pipeline'ni buzadi",
            "O'rta",
            "Yuqori",
            "Apicurio compatibility = BACKWARD. CI gate: schema PR avtomatik check.",
        ],
        [
            "Univer team Stage 2'ga vaqt topmaydi",
            "Yuqori",
            "Past (REST 100% parallel ishlaydi)",
            "Per-OTM feature flag. Hech kim majburlanmaydi. Stage 1 self-contained.",
        ],
        [
            "Kafka cluster downtime",
            "O'rta",
            "Past (outbox buffer)",
            "Outbox jadval Kafka'siz ham yoziladi. Connector restart'da backlog replay.",
        ],
        [
            "Audit DB Kafka consumer xato",
            "Past",
            "O'rta (audit gap)",
            "Audit topic 30-day retention. Manual replay imkoni.",
        ],
        [
            "Outbox publisher leader election issue (HA setup)",
            "Past",
            "O'rta",
            "ShedLock yoki Spring @SchedulerLock — bir vaqtda 1 instance.",
        ],
        [
            "Kafka producer message size limit (default 1MB)",
            "Past",
            "Past",
            "Payload size check. Large payload → S3 + reference (Claim Check pattern).",
        ],
    ],
    col_widths=[Cm(4.5), Cm(1.5), Cm(2.5), Cm(7)],
)

doc.add_page_break()

# ============================================================
# 12. KONFIGURATSIYA NAMUNALARI
# ============================================================
add_h("12. Konfiguratsiya namunalari", level=1)

add_h("12.1 docker-compose.yml (Kafka stack)", level=2)
code(
    "services:\n"
    "  kafka:\n"
    "    image: apache/kafka:3.7.0\n"
    "    environment:\n"
    "      KAFKA_NODE_ID: 1\n"
    "      KAFKA_PROCESS_ROLES: broker,controller\n"
    "      KAFKA_LISTENERS: PLAINTEXT://:9092,CONTROLLER://:9093\n"
    "      KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://kafka:9092\n"
    "      KAFKA_CONTROLLER_LISTENER_NAMES: CONTROLLER\n"
    "      KAFKA_CONTROLLER_QUORUM_VOTERS: 1@localhost:9093\n"
    "      KAFKA_LISTENER_SECURITY_PROTOCOL_MAP: PLAINTEXT:PLAINTEXT,CONTROLLER:PLAINTEXT\n"
    "      KAFKA_INTER_BROKER_LISTENER_NAME: PLAINTEXT\n"
    "      KAFKA_OFFSETS_TOPIC_REPLICATION_FACTOR: 1\n"
    "      KAFKA_AUTO_CREATE_TOPICS_ENABLE: \"false\"\n"
    "      CLUSTER_ID: 4L6g3nShT-eMCtK--X86sw\n"
    "    ports: [\"9092:9092\"]\n"
    "    volumes: [kafka-data:/var/lib/kafka/data]\n\n"
    "  apicurio:\n"
    "    image: apicurio/apicurio-registry-mem:2.5.0.Final\n"
    "    environment:\n"
    "      QUARKUS_PROFILE: prod\n"
    "    ports: [\"8888:8080\"]\n\n"
    "  kafka-ui:\n"
    "    image: provectuslabs/kafka-ui:latest\n"
    "    environment:\n"
    "      KAFKA_CLUSTERS_0_NAME: hemis-local\n"
    "      KAFKA_CLUSTERS_0_BOOTSTRAPSERVERS: kafka:9092\n"
    "      KAFKA_CLUSTERS_0_SCHEMAREGISTRY: http://apicurio:8080/apis/registry/v2\n"
    "    ports: [\"8889:8080\"]",
    label="docker-compose.yml — Kafka services:",
)

add_h("12.2 application.yml (Kafka + Outbox config)", level=2)
code(
    "hemis:\n"
    "  outbox:\n"
    "    poll-interval-ms: 5000\n"
    "    batch-size: 100\n"
    "    max-retry-count: 5\n"
    "    publishing-enabled: true\n"
    "  kafka:\n"
    "    bootstrap-servers: ${KAFKA_BROKERS:kafka:9092}\n"
    "    schema-registry-url: ${APICURIO_URL:http://apicurio:8080/apis/registry/v2}\n\n"
    "spring:\n"
    "  kafka:\n"
    "    bootstrap-servers: ${KAFKA_BROKERS:kafka:9092}\n"
    "    producer:\n"
    "      acks: all\n"
    "      compression-type: snappy\n"
    "      enable-idempotence: true\n"
    "      max-in-flight-requests-per-connection: 5\n"
    "      key-serializer: org.apache.kafka.common.serialization.StringSerializer\n"
    "      value-serializer: org.apache.kafka.common.serialization.StringSerializer\n"
    "    consumer:\n"
    "      group-id: hemis-back\n"
    "      auto-offset-reset: earliest\n"
    "      isolation-level: read_committed\n"
    "      key-deserializer: org.apache.kafka.common.serialization.StringDeserializer\n"
    "      value-deserializer: org.apache.kafka.common.serialization.StringDeserializer",
    label="application.yml:",
)

add_h("12.3 Topic creation (bootstrap)", level=2)
code(
    "# Lokal dev: topic'larni manual yaratish\n"
    "docker exec hemis-kafka /opt/kafka/bin/kafka-topics.sh \\\n"
    "    --create --topic hemis.student.events.v1 \\\n"
    "    --bootstrap-server kafka:9092 \\\n"
    "    --partitions 3 --replication-factor 1\n\n"
    "docker exec hemis-kafka /opt/kafka/bin/kafka-topics.sh \\\n"
    "    --create --topic hemis.classifier.h_position.events.v1 \\\n"
    "    --bootstrap-server kafka:9092 \\\n"
    "    --partitions 3 --replication-factor 1 \\\n"
    "    --config cleanup.policy=compact\n\n"
    "docker exec hemis-kafka /opt/kafka/bin/kafka-topics.sh \\\n"
    "    --create --topic hemis.dlq.v1 \\\n"
    "    --bootstrap-server kafka:9092 \\\n"
    "    --partitions 3 --replication-factor 1 \\\n"
    "    --config retention.ms=2592000000  # 30 day",
    label="Topic creation (one-time):",
)
p(
    "Production'da topic creation Spring `KafkaAdmin` orqali (bean), Liquibase'ga o'xshash "
    "deklarativ tarzda:"
)
code(
    "@Configuration\n"
    "public class KafkaTopicConfig {\n"
    "    @Bean\n"
    "    public NewTopic studentEvents() {\n"
    "        return TopicBuilder.name(\"hemis.student.events.v1\")\n"
    "            .partitions(3).replicas(1)\n"
    "            .build();\n"
    "    }\n"
    "}",
    label="Topic creation (Spring KafkaAdmin):",
)

doc.add_page_break()

# ============================================================
# 13. TAVSIYALAR
# ============================================================
add_h("13. Tavsiyalar (best practices)", level=1)

add_h("13.1 Birinchi domain integratsiyasi", level=2)
p(
    "Tavsiya: **ClassifierService bilan boshlang**, StudentService emas. Sabab:"
)
bullets([
    "Klassifikator domain sodda (CRUD, business logic kam)",
    "Volume past (kuniga 10-20 event), debug oson",
    "Compacted topic strategiyasi sinash imkonini beradi",
    "Back-channel use case (Hemis → Univer) bevosita namoyish qiladi",
    "Schema breaking change'lar past xavfli (klassifikator schema kam o'zgaradi)",
])

add_h("13.2 Pilot kerakmi?", level=2)
p(
    "Hozirgi qaror: pilot bosqich olib tashlandi. Sabab — production'da hali real foydalanuvchi yo'q. "
    "Lekin Stage 2 (224 OTM Kafka producer migration) keyin kelganda pilot foydali bo'lishi mumkin:"
)
bullets([
    "5-10 ta volonter OTM (TashDTU, TUIT katta + 8 kichik)",
    "1 oy validation: Kafka path latency vs REST path latency",
    "Success metric: Kafka latency ≤ REST latency, 0 ta data loss",
    "Failure path: Kafka path silently revert REST'ga (per-OTM feature flag)",
])

add_h("13.3 Production hardening (kelajakda)", level=2)
bullets([
    "Kafka cluster: 3 broker, separate disks, ACL + SASL/SCRAM auth",
    "Apicurio: PostgreSQL backed (dev'dan farqli), HA",
    "Outbox jadval cleanup job (90 day older events archived to S3)",
    "Debezium CDC qo'shish (Stage 2/3 da, volume isbotlanganda)",
    "Multi-region replication (MirrorMaker 2)",
    "Schema Registry CI gate (GitHub Actions yoki Gitlab CI)",
    "Grafana dashboard: outbox lag, Kafka producer/consumer lag, DLQ count",
    "Alert: outbox unpublished count > 1000, DLQ count > 0, consumer lag > 10s",
])

add_h("13.4 Common antipattern'lar", level=2)
table(
    ["Antipattern", "Nima'da xato", "To'g'risi"],
    [
        ["Dual-write (DB + Kafka)", "Race condition — biri muvaffaqiyatli, biri yo'q", "Outbox pattern (atomic INSERT)"],
        ["Kafka send INSIDE @Transactional", "Kafka send sync — TX uzayadi, lock contention", "Outbox + background publisher"],
        ["Topic per entity (kichik)", "Kafka 1M+ topic limit emas, lekin ops murakkab", "Topic per aggregate (domain group)"],
        ["No partition key", "Round-robin distribution → ordering yo'q", "key = aggregate_id"],
        ["Auto-create topic", "Schema bootstrap chalkash, reproducibility past", "Explicit topic creation (KafkaAdmin)"],
        ["Default consumer group", "Multiple service same group → unintended sharing", "Explicit group-id per service"],
        ["No DLQ", "Failed event loop → infinite retry queue", "DLQ topic + manual reprocess"],
    ],
    col_widths=[Cm(4), Cm(5), Cm(6)],
)

doc.add_page_break()

# ============================================================
# 14. OPEN SAVOLLAR
# ============================================================
add_h("14. Open savollar (qaror talab qiladi)", level=1)

p("Hujjat o'qib chiqilgandan so'ng quyidagi 6 ta savol bo'yicha qaror talab qilinadi:")

numbered([
    "Modul taqsimoti maqulmi? (domain: outbox entity + repo, service: publisher, app: config)",
    "Topic naming format: hemis.student.events.v1 — yoki sodda student.v1?",
    "Outbox publisher: scheduled poll (5s) yoki Debezium CDC darhol (sub-second)?",
    "Birinchi domain: ClassifierService (tavsiya etilgan) yoki StudentService?",
    "Service ↔ Outbox bog'lanish strategy:\n"
    "    a) Eksplisit chaqiruv: outbox.publish(...) (tavsiya etilgan, debug oson)\n"
    "    b) AOP aspect (avtomatik, lekin \"magical\")\n"
    "    c) ApplicationEventPublisher → Kafka (Spring native)",
    "Implementation tartibi: A) men 10 qadamni ketma-ket, B) har qadamdan keyin men sizdan tasdiq olaman?",
])

doc.add_page_break()

# ============================================================
# 15. XULOSA VA KEYINGI QADAMLAR
# ============================================================
add_h("15. Xulosa va keyingi qadamlar", level=1)

add_h("15.1 Asosiy fikrlar", level=2)
bullets([
    "Hozirgi REST sync 8 ta muammo bilan — eng katta xavf data loss",
    "8 ta alternativ baholandi, Kafka + Outbox Pattern eng yaxshi mos keladi",
    "Greenfield holat — pilot olib tashlandi, to'g'ri arxitektura bilan boshlanadi",
    "REST endpoint'lar 12+ oy parallel saqlanadi — backward compat 100%",
    "224 OTM o'z tezligida Kafka producer'ga o'tadi (Stage 2)",
    "Stage 1 — 4-6 hafta ish, ~3 soat code + test/verify",
])

add_h("15.2 Keyingi qadamlar (tartib)", level=2)
numbered([
    "Sizning qaror ma'qullashingiz (14-bo'limdagi 6 ta savol)",
    "ADR-0007 yangilanishi (qabul qilingan qaror bo'yicha)",
    "10 ta atomik qadam ketma-ket bajariladi (har qadamda build verify)",
    "Step 11+: birinchi domain integratsiyasi (ClassifierService tavsiya)",
    "Topic creation + Schema Registry sample",
    "Audit consumer (hemis_audit DB ulanishi)",
    "Stage 2 (224 OTM migration) — alohida sprint, OTM team tayyorgarligi bilan",
])

add_h("15.3 References", level=2)
bullets([
    "Chris Richardson, Microservices Patterns (2018) — Outbox, Saga, CQRS",
    "Ben Stopford, Designing Event-Driven Systems (2018) — Kafka topology",
    "Martin Kleppmann, Designing Data-Intensive Applications (2017)",
    "Sam Newman, Building Microservices, 2nd ed. (2021) — Strangler Fig",
    "Apache Kafka KRaft mode docs — kafka.apache.org/documentation",
    "Apicurio Schema Registry — apicur.io/registry",
    "Stripe Idempotency-Key API design — stripe.com/blog/idempotency",
    "Confluent Schema Compatibility — docs.confluent.io",
    "ADR-0003 (Audit DB isolation)",
    "ADR-0004 (api-university module)",
    "ADR-0005 (OAuth client_credentials)",
    "ADR-0006 (h_* classifier prefix)",
    "ADR-0007 (Sync Architecture — Kafka-first Approach)",
    "docs/UNIVER_CONTRACT.md — 224 OTM frozen API contract",
    "docs/UNIVER_ENDPOINT_AUDIT.md — per-endpoint audit",
])

# Footer
doc.add_paragraph()
hrule()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("HEMIS-back team   |   Sync Architecture Report   |   2026-05-06")
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ============================================================
# Save
# ============================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"DOCX yaratildi: {OUT}")
print(f"Hajm: {OUT.stat().st_size} bayt")
